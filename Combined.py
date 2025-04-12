import streamlit as st
import requests
import base64
import json
from supabase import create_client
from datetime import datetime, timedelta
import re
import time
import google.generativeai as genai
import docx
from docx.shared import Pt
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import os
import pathlib
import uuid
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set page config
st.set_page_config(
    page_title="Learning Observer",
    layout="wide",
    page_icon="📝"
)


# Initialize Supabase client with enhanced error handling
@st.cache_resource
def init_supabase():
    try:
        SUPABASE_URL = st.secrets.get("SUPABASE_URL")
        SUPABASE_KEY = st.secrets.get("SUPABASE_KEY")

        if not SUPABASE_URL or not SUPABASE_KEY:
            raise ValueError("Supabase URL or KEY not found in secrets")

        client = create_client(SUPABASE_URL, SUPABASE_KEY)

        # Test connection
        try:
            test = client.table('users').select("count", count="exact").execute()
            logger.info(f"Supabase connected successfully. Found {test.count} users.")
        except Exception as test_error:
            logger.error(f"Supabase connection test failed: {str(test_error)}")
            st.error("Failed to connect to database. Please check your connection settings.")
            raise

        return client

    except Exception as e:
        logger.error(f"Supabase initialization failed: {str(e)}")
        st.error("Database initialization failed. Please check your configuration.")
        raise


supabase = init_supabase()

# Configure Google AI API
genai.configure(api_key=st.secrets.get("GOOGLE_API_KEY"))

# Set up AssemblyAI API key
assemblyai_key = st.secrets.get("ASSEMBLYAI_API_KEY", "")


class ObservationExtractor:
    def __init__(self):
        self.ocr_api_key = st.secrets.get("OCR_API_KEY")
        self.groq_api_key = st.secrets.get("GROQ_API_KEY")
        self.gemini_api_key = st.secrets.get("GOOGLE_API_KEY")

    def image_to_base64(self, image_file):
        """Convert image file to base64 string"""
        return base64.b64encode(image_file.read()).decode('utf-8')

    def extract_text_with_ocr(self, image_file):
        """Extract text from image using OCR.space API"""
        try:
            # Get file extension
            file_type = image_file.name.split('.')[-1].lower()
            if file_type == 'jpeg':
                file_type = 'jpg'

            # Convert image to base64
            base64_image = self.image_to_base64(image_file)
            base64_image_with_prefix = f"data:image/{file_type};base64,{base64_image}"

            # Prepare request payload
            payload = {
                'apikey': self.ocr_api_key,
                'language': 'eng',
                'isOverlayRequired': False,
                'iscreatesearchablepdf': False,
                'issearchablepdfhidetextlayer': False,
                'OCREngine': 2,  # Better for handwriting
                'detectOrientation': True,
                'scale': True,
                'base64Image': base64_image_with_prefix
            }

            # Send request to OCR API
            response = requests.post(
                'https://api.ocr.space/parse/image',
                data=payload,
                headers={'apikey': self.ocr_api_key}
            )

            response.raise_for_status()
            data = response.json()

            # Process response
            if not data.get('ParsedResults') or len(data['ParsedResults']) == 0:
                error_msg = data.get('ErrorMessage', 'No parsed results returned')
                raise Exception(f"OCR Error: {error_msg}")

            parsed_result = data['ParsedResults'][0]
            if parsed_result.get('ErrorMessage'):
                raise Exception(f"OCR Error: {parsed_result['ErrorMessage']}")

            extracted_text = parsed_result['ParsedText']

            if not extracted_text or not extracted_text.strip():
                raise Exception("No text was detected in the image")

            return extracted_text

        except Exception as e:
            st.error(f"OCR Error: {str(e)}")
            raise

    def process_with_groq(self, extracted_text):
        """Process extracted text with Groq AI"""
        try:
            # Original detailed prompt
            system_prompt = """You are an AI assistant for a learning observation system. Extract and structure information from the provided observation sheet text.

The observation sheets typically have the following structure:
- Title (usually "The Observer")
- Student information (Name, Roll Number/ID)
- Date and Time information
- Core Observation Section with time slots
- Teaching content for each time slot
- Learning details (what was learned, tools used, etc.)

Format your response as JSON with the following structure:
{
  "studentName": "Student's name if available, otherwise use the title of the observation",
  "studentId": "Student ID or Roll Number",
  "className": "Class name or subject being taught",
  "date": "Date of observation",
  "observations": "Detailed description of what was learned",
  "strengths": ["List of strengths observed in the student"],
  "areasOfDevelopment": ["List of areas where the student needs improvement"],
  "recommendations": ["List of recommended actions for improvement"]
}

For observations, provide full detailed descriptions like:
"The student learned how to make maggi from their mom through in-person mode, including all steps from boiling water to adding spices"

Be creative in extracting information based on context."""

            # Send request to Groq API
            response = requests.post(
                'https://api.groq.com/openai/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {self.groq_api_key}',
                    'Content-Type': 'application/json'
                },
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {
                            "role": "system",
                            "content": system_prompt
                        },
                        {
                            "role": "user",
                            "content": f"Extract and structure the following text from an observation sheet: {extracted_text}"
                        }
                    ],
                    "temperature": 0.2,
                    "response_format": {"type": "json_object"}
                }
            )

            response.raise_for_status()
            data = response.json()

            # Extract the JSON content
            ai_response = data['choices'][0]['message']['content']
            return json.loads(ai_response)

        except Exception as e:
            st.error(f"Groq API Error: {str(e)}")
            raise

    def transcribe_with_assemblyai(self, audio_file):
        """Transcribe audio using AssemblyAI API"""
        if not assemblyai_key:
            return "Error: AssemblyAI API key is not configured. Please add it to your secrets."

        # Set up the API headers
        headers = {
            "authorization": assemblyai_key,
            "content-type": "application/json"
        }

        # Upload the audio file
        try:
            st.write("Uploading audio file...")
            upload_response = requests.post(
                "https://api.assemblyai.com/v2/upload",
                headers={"authorization": assemblyai_key},
                data=audio_file.getvalue()
            )

            if upload_response.status_code != 200:
                return f"Error uploading audio: {upload_response.text}"

            upload_url = upload_response.json()["upload_url"]

            # Request transcription
            st.write("Processing transcription...")
            transcript_request = {
                "audio_url": upload_url,
                "language_code": "en"
            }

            transcript_response = requests.post(
                "https://api.assemblyai.com/v2/transcript",
                json=transcript_request,
                headers=headers
            )

            if transcript_response.status_code != 200:
                return f"Error requesting transcription: {transcript_response.text}"

            transcript_id = transcript_response.json()["id"]

            # Poll for completion
            status = "processing"
            progress_bar = st.progress(0)
            while status != "completed" and status != "error":
                polling_response = requests.get(
                    f"https://api.assemblyai.com/v2/transcript/{transcript_id}",
                    headers=headers
                )

                if polling_response.status_code != 200:
                    return f"Error checking transcription status: {polling_response.text}"

                polling_data = polling_response.json()
                status = polling_data["status"]

                if status == "completed":
                    progress_bar.progress(100)
                    return polling_data["text"]
                elif status == "error":
                    return f"Transcription error: {polling_data.get('error', 'Unknown error')}"

                # Update progress
                progress = polling_data.get("percent_done", 0)
                if progress:
                    progress_bar.progress(progress / 100.0)
                time.sleep(2)

            return "Error: Transcription timed out or failed."
        except Exception as e:
            return f"Error during transcription: {str(e)}"

    def generate_report_from_text(self, text_content, user_info):
        """Generate a structured report from text using Google Gemini"""
        prompt = f"""
        Based on this text from a student observation, create a detailed observer report following the Observer Report format.

        TEXT CONTENT:
        {text_content}

        FORMAT REQUIREMENTS:

        1. Daily Activities Overview - Extract and categorize the student's daily activities into:
           - Morning activities
           - Afternoon activities
           - Evening activities
           - Night activities (if mentioned)

        2. Learning Moments & Reflections - Identify:
           - New skills or knowledge the student gained
           - Interesting observations or experiences
           - Any self-reflection shared

        3. Thinking Process - Assess:
           - Approach to new information (curious/skeptical/analytical/accepting)
           - Logical thinking (strong/moderate/needs improvement)
           - Problem-solving skills (effective/developing/needs guidance)
           - Creativity and imagination (high/moderate/low)
           - Decision-making style (confident/hesitant/experimental)
           - Any unique perspectives or ideas

        4. Communication Skills & Thought Clarity - Evaluate:
           - Confidence level (low/medium/high)
           - Clarity of thought (clear/slightly clear/confused)
           - Participation & engagement (active/moderate/passive)
           - Sequence of explanation (well-structured/partially structured/unstructured)

        5. General Behavior & Awareness - Note:
           - Behavior (polite/calm/energetic/distracted)
           - General awareness (aware/partially aware/unaware)

        6. Observer's Comments - Add any relevant observations

        7. Summary for Parents - Write a brief paragraph summarizing the session

        Use the exact section titles and format as above. For items that cannot be determined from the text, use "Not enough information" rather than making assumptions.
        """

        try:
            # Configure the model - using Gemini Pro for most comprehensive responses
            model = genai.GenerativeModel('gemini-1.5-pro-002')

            # Generate content with Gemini
            response = model.generate_content([
                {"role": "user", "parts": [{"text": prompt}]}
            ])

            # Extract the content from the response
            report_content = response.text

            # Add user information to the report
            complete_report = f"""Date: {user_info['session_date']}
    Student Name: {user_info['student_name']}
    Observer Name: {user_info['observer_name']}
    Session Duration: {user_info['session_start']} - {user_info['session_end']}

    {report_content}

    Name of Observer: {user_info['observer_name']}
    """
            return complete_report
        except Exception as e:
            return f"Error generating report: {str(e)}"

    def create_word_document(self, report_content):
        """Create a Word document from the report content with proper formatting"""
        doc = docx.Document()

        # Add title
        title = doc.add_heading('The Observer Report', 0)

        # Clean up markdown formatting
        report_content = report_content.replace('**', '')

        # Add report content, parsing the sections
        lines = report_content.split('\n')
        section_pattern = re.compile(r'^\d+\.\s+(.+)')
        subheading_pattern = re.compile(r'^\*\s*(.+):\*\s*(.+)')
        list_item_pattern = re.compile(r'^\*\s+(.+)')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Header information (Date, Name, etc.)
            if line.startswith(('Date:', 'Student Name:', 'Observer Name:', 'Session Duration:', 'Name of Observer:')):
                p = doc.add_paragraph()
                p.add_run(line).bold = True

            # Section heading (e.g., "1. Daily Activities Overview")
            elif section_match := section_pattern.match(line):
                doc.add_heading(line, level=1)

            # Subheading with bold prefix (e.g., "* Morning activities: Woke up early")
            elif subheading_match := subheading_pattern.match(line):
                p = doc.add_paragraph()
                prefix = subheading_match.group(1)
                content = subheading_match.group(2)
                p.add_run(f"{prefix}: ").bold = True
                p.add_run(content)

            # List item
            elif list_match := list_item_pattern.match(line):
                content = list_match.group(1)
                p = doc.add_paragraph(content, style='List Bullet')

            # Regular paragraph
            else:
                doc.add_paragraph(line)

        # Save to a BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        return docx_bytes

    def send_email(self, recipient_email, subject, message):
        """Send email with the observation report"""
        sender_email = "parth.workforai@gmail.com"
        sender_password = st.secrets.get("EMAIL_PASSWORD")  # Add this to your secrets.toml

        smtp_server = "smtp.gmail.com"
        smtp_port = 587

        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = recipient_email
        msg["Subject"] = subject
        msg.attach(MIMEText(message, "html"))

        try:
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
            return True, f"Email sent to {recipient_email}"
        except smtplib.SMTPAuthenticationError:
            return False, "Error: Authentication failed. Check your email and password."
        except smtplib.SMTPException as e:
            return False, f"Error: Failed to send email. {e}"
        finally:
            server.quit()


def admin_dashboard():
    st.title("Admin Dashboard")

    tabs = st.tabs(["Mappings", "Activity", "Users"])

    with tabs[0]:
        st.subheader("Observer-Child Mappings")
        try:
            mappings = supabase.table('observer_child_mappings').select("*").execute().data
            if mappings:
                # Add a column for delete action buttons
                for i, mapping in enumerate(mappings):
                    col1, col2, col3 = st.columns([3, 3, 1])
                    with col1:
                        st.write(f"Observer ID: {mapping['observer_id']}")
                    with col2:
                        st.write(f"Child ID: {mapping['child_id']}")
                    with col3:
                        if st.button("Delete", key=f"delete_{mapping['id']}"):
                            supabase.table('observer_child_mappings').delete().eq('id', mapping['id']).execute()
                            st.success("Mapping deleted successfully!")
                            st.rerun()
            else:
                st.info("No mappings found")

            with st.expander("Add New Mapping"):
                with st.form("add_mapping"):
                    obs_id = st.text_input("Observer ID")
                    child_id = st.text_input("Child ID")
                    if st.form_submit_button("Add"):
                        if obs_id and child_id:
                            supabase.table('observer_child_mappings').insert({
                                "observer_id": obs_id,
                                "child_id": child_id
                            }).execute()
                            st.rerun()
        except Exception as e:
            st.error(f"Database error: {str(e)}")

    with tabs[1]:
        st.subheader("Activity Logs")
        try:
            logs = supabase.table('observer_activity_logs').select("*").execute().data
            if logs:
                st.dataframe(logs)
            else:
                st.info("No activity logs")
        except Exception as e:
            st.error(f"Database error: {str(e)}")

    with tabs[2]:
        st.subheader("User Management")
        try:
            users = supabase.table('users').select("*").execute().data
            if users:
                st.write("Registered Users:")
                for user in users:
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"Name: {user.get('name', 'N/A')} ({user['email']})")
                    with col2:
                        st.write(f"Role: {user['role']}")
                    with col3:
                        if st.button("Delete", key=f"delete_{user['id']}"):
                            supabase.table('users').delete().eq('id', user['id']).execute()
                            st.rerun()
            else:
                st.info("No users found")
        except Exception as e:
            st.error(f"Database error: {str(e)}")


# Parent Dashboard
def parent_dashboard(user_id):
    st.title(f"Parent Portal")

    try:
        # Get the child ID for this parent
        parent_data = supabase.table('users').select("child_id").eq("id", user_id).execute().data
        if not parent_data or not parent_data[0].get('child_id'):
            st.warning("No child assigned to your account. Please contact admin.")
            return

        child_id = parent_data[0]['child_id']

        # Get observer ID
        mapping = supabase.table('observer_child_mappings').select("observer_id").eq("child_id",
                                                                                     child_id).execute().data
        if not mapping:
            st.warning("No observer assigned to your child")
            return

        observer_id = mapping[0]['observer_id']

        # Get observer name
        observer_data = supabase.table('users').select("name").eq("id", observer_id).execute().data
        observer_name = observer_data[0]['name'] if observer_data else observer_id

        # Get activity logs
        logs = supabase.table('observer_activity_logs').select("*").eq("child_id", child_id).execute().data

        # Calculate metrics
        total_logins = len([log for log in logs if log['action'] == 'login'])
        total_time = sum(log.get('duration_minutes', 0) for log in logs)
        reports_submitted = len([log for log in logs if log.get('report_id')])

        # Display info
        cols = st.columns(3)
        cols[0].metric("Assigned Observer", observer_name)
        cols[1].metric("Total Logins", total_logins)
        cols[2].metric("Total Observation Time (mins)", total_time)

        st.subheader("Recent Reports")
        reports = supabase.table('observations').select("*").eq("student_id", child_id).execute().data
        if reports:
            for report in reports:
                with st.expander(f"Report from {report.get('date', 'unknown date')}"):
                    st.write(f"**Student:** {report.get('student_name', 'N/A')}")
                    st.write(f"**Date:** {report.get('date', 'N/A')}")
                    st.write(f"**Observations:**")
                    st.write(report.get('observations', 'No observations recorded'))

                    # Display strengths, areas of development, and recommendations if available
                    if report.get('strengths'):
                        st.write("**Strengths:**")
                        strengths = json.loads(report['strengths'])
                        for strength in strengths:
                            st.write(f"- {strength}")

                    if report.get('areas_of_development'):
                        st.write("**Areas for Development:**")
                        areas = json.loads(report['areas_of_development'])
                        for area in areas:
                            st.write(f"- {area}")

                    if report.get('recommendations'):
                        st.write("**Recommendations:**")
                        recs = json.loads(report['recommendations'])
                        for rec in recs:
                            st.write(f"- {rec}")
        else:
            st.info("No reports available yet")

    except Exception as e:
        st.error(f"Database error: {str(e)}")


# Main App
def main():
    extractor = ObservationExtractor()

    # Session State Initialization
    if 'auth' not in st.session_state:
        st.session_state.auth = {
            'logged_in': False,
            'role': None,
            'user_id': None,
            'email': None,
            'name': None
        }
    if 'user_info' not in st.session_state:
        st.session_state.user_info = {
            'student_name': '',
            'observer_name': '',
            'session_date': datetime.now().strftime('%d/%m/%Y'),
            'session_start': '',
            'session_end': ''
        }
    if 'audio_transcription' not in st.session_state:
        st.session_state.audio_transcription = ""
    if 'report_generated' not in st.session_state:
        st.session_state.report_generated = None
    if 'show_edit_transcript' not in st.session_state:
        st.session_state.show_edit_transcript = False
    if 'processing_mode' not in st.session_state:
        st.session_state.processing_mode = None
    if 'show_register' not in st.session_state:
        st.session_state.show_register = False
    if 'admin_initial_login' not in st.session_state:
        st.session_state.admin_initial_login = True

    # Admin credentials
    ADMIN_CREDS = {
        "username": st.secrets.get("ADMIN_USER", "admin"),
        "password": st.secrets.get("ADMIN_PASS", "hello")
    }

    # Login/Registration Page
    if not st.session_state.auth['logged_in']:
        st.title("Learning Observer Login")

        if st.session_state.show_register:
            # Registration Form with improved validation
            with st.form("register_form"):
                st.subheader("Create New Account")
                name = st.text_input("Full Name", max_chars=100)
                email = st.text_input("Email", max_chars=100).strip().lower()
                role = st.selectbox("Role", ["Observer", "Parent"])
                password = st.text_input("Password", type="password", max_chars=100)
                confirm_password = st.text_input("Confirm Password", type="password", max_chars=100)

                child_id = None
                if role == "Parent":
                    child_id = st.text_input("Child ID (provided by your observer)", max_chars=50)

                submitted = st.form_submit_button("Register")

                if submitted:
                    if not all([name, email, password, confirm_password]):
                        st.error("Please fill in all required fields")
                    elif password != confirm_password:
                        st.error("Passwords do not match!")
                    elif len(password) < 8:
                        st.error("Password must be at least 8 characters")
                    else:
                        try:
                            # Check if email exists
                            existing_response = supabase.table('users').select("email").eq("email", email).execute()

                            if existing_response.data:
                                st.error("Email already registered. Please login instead.")
                            else:
                                # Prepare user data with explicit column names
                                user_data = {
                                    "id": str(uuid.uuid4()),
                                    "email": email,
                                    "name": name,
                                    "password": password,  # Note: In production, hash this password!
                                    "role": role
                                }

                                if role == "Parent" and child_id:
                                    user_data["child_id"] = child_id.strip()

                                # Insert with error handling
                                insert_response = supabase.table('users').insert(user_data).execute()

                                if insert_response.data:
                                    st.success("Registration successful! Please login.")
                                    st.session_state.show_register = False
                                else:
                                    st.error("Registration failed. Please try again.")

                        except Exception as e:
                            logger.error(f"Registration error: {str(e)}")
                            st.error(f"Registration failed: {str(e)}")

            if st.button("Back to Login"):
                st.session_state.show_register = False
        else:
            # Login Form with improved error handling
            with st.form("login_form"):
                email = st.text_input("Email", max_chars=100).strip().lower()
                password = st.text_input("Password", type="password", max_chars=100)

                submitted = st.form_submit_button("Login")

                if submitted:
                    try:
                        # Check admin login first
                        if email == ADMIN_CREDS["username"] and password == ADMIN_CREDS["password"]:
                            st.session_state.auth = {
                                'logged_in': True,
                                'role': 'Admin',
                                'user_id': 'admin',
                                'email': email,
                                'name': 'Admin'
                            }
                            st.rerun()

                        # Check regular user login
                        user_response = supabase.table('users').select("*").eq("email", email).eq("password",
                                                                                                  password).execute()

                        if user_response.data:
                            user = user_response.data[0]
                            st.session_state.auth = {
                                'logged_in': True,
                                'role': user['role'],
                                'user_id': user['id'],
                                'email': user['email'],
                                'name': user.get('name', 'User')
                            }
                            st.rerun()
                        else:
                            st.error("Invalid email or password")

                    except Exception as e:
                        logger.error(f"Login error: {str(e)}")
                        st.error(f"Login failed: {str(e)}")

            st.write("Don't have an account?")
            if st.button("Register Here"):
                st.session_state.show_register = True

        return

    # Logout Button (common)
    def logout_button():
        if st.button("Logout"):
            if st.session_state.auth['role'] == "Observer":
                supabase.table('observer_activity_logs').insert({
                    "observer_id": st.session_state.auth['user_id'],
                    "child_id": "N/A",
                    "action": "logout",
                    "duration_minutes": 0
                }).execute()
            st.session_state.auth = {'logged_in': False, 'role': None, 'user_id': None}
            st.rerun()

    # Admin Dashboard - Modified to check for initial login
    if st.session_state.auth['role'] == 'Admin':
        if st.session_state.admin_initial_login:
            # Show a welcome or intermediate page instead of the full admin dashboard
            st.title("Welcome, Admin")
            st.write("You are logged in as an administrator.")

            # Add a button to proceed to the dashboard
            if st.button("Access Admin Dashboard"):
                st.session_state.admin_initial_login = False
                st.rerun()

            logout_button()
            return
        else:
            # Show the regular admin dashboard
            admin_dashboard()
            logout_button()
            return

    # Parent Dashboard
    if st.session_state.auth['role'] == 'Parent':
        parent_dashboard(st.session_state.auth['user_id'])
        logout_button()
        return

    # Observer Dashboard
    st.title(f"Observer Dashboard - {st.session_state.auth['name']}")

    # Log login activity
    supabase.table('observer_activity_logs').insert({
        "observer_id": st.session_state.auth['user_id'],
        "child_id": "N/A",
        "action": "login",
        "duration_minutes": 0
    }).execute()

    logout_button()

    # Sidebar for user information
    with st.sidebar:
        st.subheader("Session Information")
        st.session_state.user_info['student_name'] = st.text_input("Student Name:",
                                                                   value=st.session_state.user_info['student_name'])
        st.session_state.user_info['observer_name'] = st.text_input("Observer Name:",
                                                                    value=st.session_state.user_info[
                                                                              'observer_name'] or st.session_state.auth.get(
                                                                        'name', ''))
        st.session_state.user_info['session_date'] = st.date_input("Session Date:").strftime('%d/%m/%Y')
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.user_info['session_start'] = st.time_input("Start Time:").strftime('%H:%M')
        with col2:
            st.session_state.user_info['session_end'] = st.time_input("End Time:").strftime('%H:%M')

    # Choose processing mode
    st.subheader("Select Processing Mode")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("OCR Mode (Image Upload)"):
            st.session_state.processing_mode = "ocr"
            st.session_state.audio_transcription = ""
            st.session_state.report_generated = None
    with col2:
        if st.button("Audio Mode (Recording Upload)"):
            st.session_state.processing_mode = "audio"
            st.session_state.audio_transcription = ""
            st.session_state.report_generated = None

    # OCR Processing
    if st.session_state.processing_mode == "ocr":
        st.info("OCR Mode: Upload an image of an observation sheet")
        uploaded_file = st.file_uploader("Upload Observation Sheet", type=["jpg", "jpeg", "png"])
        if uploaded_file and st.button("Process Observation"):
            with st.spinner("Processing..."):
                try:
                    extracted_text = extractor.extract_text_with_ocr(uploaded_file)
                    structured_data = extractor.process_with_groq(extracted_text)
                    observations_text = structured_data.get("observations", "")
                    if observations_text:
                        report = extractor.generate_report_from_text(observations_text, st.session_state.user_info)
                        st.session_state.report_generated = report

                        # Get child ID from mappings if available
                        child_id = structured_data.get("studentId", "")
                        if not child_id:
                            # Try to find child by name
                            child_data = supabase.table('users').select("id").ilike("name",
                                                                                    f"%{structured_data.get('studentName', '')}%").execute().data
                            if child_data:
                                child_id = child_data[0]['id']

                        supabase.table('observations').insert({
                            "username": st.session_state.auth['user_id'],
                            "student_name": structured_data.get("studentName", ""),
                            "student_id": child_id,
                            "class_name": structured_data.get("className", ""),
                            "date": structured_data.get("date", ""),
                            "observations": observations_text,
                            "strengths": json.dumps(structured_data.get("strengths", [])),
                            "areas_of_development": json.dumps(structured_data.get("areasOfDevelopment", [])),
                            "recommendations": json.dumps(structured_data.get("recommendations", [])),
                            "timestamp": datetime.now().isoformat(),
                            "filename": uploaded_file.name,
                            "full_data": json.dumps(structured_data)
                        }).execute()
                        st.success("Data processed and saved successfully!")
                    else:
                        st.error("No observations found in the extracted data")
                except Exception as e:
                    st.error(f"Processing error: {str(e)}")

    # Audio Processing
    elif st.session_state.processing_mode == "audio":
        st.info("Audio Mode: Upload an audio recording of an observation session")
        uploaded_file = st.file_uploader("Choose an audio file",
                                         type=["wav", "mp3", "m4a", "mpeg", "ogg", "flac", "aac", "wma", "aiff"])
        if uploaded_file and st.button("Process & Generate Report"):
            if not assemblyai_key:
                st.error("AssemblyAI API key is missing.")
            else:
                with st.spinner("Step 1/2: Transcribing audio..."):
                    transcript = extractor.transcribe_with_assemblyai(uploaded_file)
                    st.session_state.audio_transcription = transcript
                with st.spinner("Step 2/2: Generating report..."):
                    report = extractor.generate_report_from_text(transcript, st.session_state.user_info)
                    st.session_state.report_generated = report
                    supabase.table('observations').insert({
                        "username": st.session_state.auth['user_id'],
                        "student_name": st.session_state.user_info['student_name'],
                        "student_id": "",
                        "class_name": "",
                        "date": st.session_state.user_info['session_date'],
                        "observations": transcript,
                        "strengths": json.dumps([]),
                        "areas_of_development": json.dumps([]),
                        "recommendations": json.dumps([]),
                        "timestamp": datetime.now().isoformat(),
                        "filename": uploaded_file.name,
                        "full_data": json.dumps({"transcript": transcript, "report": report})
                    }).execute()

    # Transcript Editor
    if st.session_state.audio_transcription:
        if st.button("Edit Transcription" if not st.session_state.show_edit_transcript else "Hide Editor"):
            st.session_state.show_edit_transcript = not st.session_state.show_edit_transcript
        if st.session_state.show_edit_transcript:
            st.subheader("Edit Transcription")
            edited = st.text_area("Edit transcript below:", value=st.session_state.audio_transcription, height=200)
            if edited != st.session_state.audio_transcription:
                st.session_state.audio_transcription = edited
            if st.button("Regenerate Report with Edited Transcript"):
                with st.spinner("Regenerating report..."):
                    report = extractor.generate_report_from_text(edited, st.session_state.user_info)
                    st.session_state.report_generated = report

    # Report Display and Download
    if st.session_state.report_generated:
        st.subheader("Generated Report")
        st.markdown(st.session_state.report_generated)
        docx_file = extractor.create_word_document(st.session_state.report_generated)
        student = st.session_state.user_info['student_name'].replace(" ", "_")
        date = st.session_state.user_info['session_date'].replace("/", "-")
        filename = f"Observer_Report_{student}_{date}.docx"
        st.download_button("Download as Word Document", data=docx_file, file_name=filename,
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.subheader("Email Report")
        with st.form("email_form"):
            to_email = st.text_input("Recipient Email", value="parent@example.com")
            subject = st.text_input("Subject",
                                    value=f"Observer Report for {st.session_state.user_info['student_name']}")
            submitted = st.form_submit_button("Send Email")
            if submitted:
                success, message = extractor.send_email(to_email, subject, st.session_state.report_generated)
                if success:
                    st.success(message)
                else:
                    st.error(message)


if __name__ == "__main__":
    main()
